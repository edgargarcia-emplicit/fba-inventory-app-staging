"""
docx_parser.py — Parses Amazon copy-template .docx files into per-ASIN
content maps (title, bullets 1-5, description, keywords).

Ported from the WordPress plugin's FLM_QA_DocxParser + FLM_CrossRef
description/keyword extraction, using python-docx for text extraction
instead of raw XML regex (more robust — python-docx correctly handles
runs within a paragraph, which the PHP fallback approach worked around
with its own regex).

Supported patterns (unchanged from the original):
  - Parent title:   "Title text (B0XXXXXXXX)"
  - Child titles:   "Title text - Variant (B0XXXXXXXX)"
  - Bullets for specific ASINs: "Bullet text (B0XXX,B0YYY)"
  - Bullets for all ASINs: "Bullet text" (no ASIN suffix)
  - TEST TITLE / CURRENT TITLE sections, when both are present the
    caller must choose which to use (title_mode='test'|'current').
"""

import re

import docx

ASIN_RE = re.compile(r"^B[A-Z0-9]{9}$")
TRAILING_PARENS_RE = re.compile(r"\(([^)]+)\)\s*$")
ANY_ASIN_PARENS_RE = re.compile(r"\(([B0][A-Z0-9,\s]+)\)")
STRIP_ASIN_SUFFIX_RE = re.compile(r"\s*\([B0][A-Z0-9,\s]+\)\s*$")

SECTION_KEYWORDS = [
    "PARENT TITLE", "CHILD TITLE", "BULLET POINT",
    "BACKEND", "DESCRIPTION", "SEO", "ADS HEADLINE",
    "KEYWORDS", "HIGH-VALUE", "TEST TITLE", "CURRENT TITLE",
    "METADATA", "ALT TEXT", "BACKEND SEARCH",
    "ABOUT THE BRAND", "BRAND STORY", "BRAND DESCRIPTION",
    "THUMBNAILS", "LISTING IMAGE", "MAIN IMAGE", "SECOND MAIN",
    "SUPPLEMENTAL", "VIDEO", "A+ CONTENT", "A+ DESCRIPTION",
    "ENHANCED BRAND", "EBC",
]

SEO_HEADERS = ["SEO", "BACKEND SEARCH TERMS", "BACKEND SEARCH"]
SEO_MARKERS = ["BACKEND SEARCH TERMS", "BACKEND SEARCH", "MAX. 250", "MAX. 500",
               "MAX. 1000", "PLEASE NOTE", "THE FOLLOWING KEYWORDS",
               "ADS HEADLINE", "AD HEADLINE"]
SEO_STOP_HEADERS = ["METADATA", "ALT TEXT", "ONE MORE THING", "TITLE", "BULLET",
                    "DESCRIPTION", "THUMBNAILS", "LISTING IMAGE", "MAIN IMAGE",
                    "SECOND MAIN", "SUPPLEMENTAL", "ABOUT THE BRAND", "BRAND STORY",
                    "BRAND DESCRIPTION", "VIDEO", "A+ CONTENT"]


def _iter_block_items(parent):
    """
    Yield each paragraph and table in the document, IN THE ORDER THEY ACTUALLY
    APPEAR in the body. python-docx's .paragraphs and .tables are each flat
    lists in their own document order, but reading them separately (all
    paragraphs, then all tables) loses the true interleaved order — which
    breaks section-boundary detection when a section's real content lives
    inside a table between two paragraph headers (exactly what these copy
    templates do for the EBC/description and SEO sections).
    """
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    for child in parent.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def extract_lines(file_path: str) -> list[str]:
    """Extract non-empty, whitespace-normalized text in true document reading order."""
    doc = docx.Document(file_path)
    lines = []
    for block in _iter_block_items(doc):
        if hasattr(block, "text"):  # Paragraph
            text = re.sub(r"\s+", " ", block.text).strip()
            if text:
                lines.append(text)
        else:  # Table
            for row in block.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = re.sub(r"\s+", " ", para.text).strip()
                        if text:
                            lines.append(text)
    return lines


def _is_section_header(line_upper: str) -> bool:
    return any(kw in line_upper for kw in SECTION_KEYWORDS)


def _extract_asins(line: str) -> list[str]:
    """Extract ASINs from the trailing (or any) parenthetical group in a line."""
    m = TRAILING_PARENS_RE.findall(line)
    if not m:
        m = ANY_ASIN_PARENS_RE.findall(line)
    if not m:
        return []
    candidate = m[-1]
    parts = re.split(r"[\s,]+", candidate)
    asins = []
    for part in parts:
        part = part.strip().upper()
        if ASIN_RE.match(part):
            asins.append(part)
    return asins


def _strip_asin_suffix(line: str) -> str:
    return STRIP_ASIN_SUFFIX_RE.sub("", line).strip()


def _has_section(lines: list[str], keyword: str) -> bool:
    return any(keyword.lower() in line.lower() for line in lines)


def _parse_titles(lines: list[str], title_mode: str, has_test: bool, has_current: bool) -> dict:
    asins = {}
    parent_labels = ["PARENT TITLE"]
    child_labels = ["CHILD TITLE"]

    if has_test and has_current:
        wanted = "TEST TITLE" if title_mode == "test" else "CURRENT TITLE"
        parent_labels.append(wanted)
        child_labels.append(wanted)

    in_title_section = False
    skip_mode = False

    for line in lines:
        line_upper = line.upper()

        if _is_section_header(line_upper):
            in_title_section = False
            skip_mode = False
            for label in parent_labels + child_labels:
                if label in line_upper:
                    in_title_section = True
                    break
            if has_test and has_current:
                unwanted = "CURRENT TITLE" if title_mode == "test" else "TEST TITLE"
                if unwanted in line_upper:
                    in_title_section = False
                    skip_mode = True
            if "BULLET" in line_upper:
                in_title_section = False
            continue

        if not in_title_section or skip_mode:
            continue

        line_asins = _extract_asins(line)
        if not line_asins:
            continue

        title_text = _strip_asin_suffix(line)
        for asin in line_asins:
            if asin not in asins:
                asins[asin] = {
                    "asin": asin, "sku": "", "title": title_text, "brand": "",
                    "bullet_1": "", "bullet_2": "", "bullet_3": "", "bullet_4": "", "bullet_5": "",
                }
    return asins


def _parse_bullet_line(line: str) -> dict:
    line = line.strip()
    asins = _extract_asins(line)
    text = _strip_asin_suffix(line)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    return {"text": text.strip(), "asins": asins}


def _parse_bullets(lines: list[str]) -> list[dict]:
    bullets = []
    in_bullet_section = False
    buffer = ""

    for line in lines:
        line_upper = line.upper()

        if "BULLET POINT" in line_upper and _is_section_header(line_upper):
            in_bullet_section = True
            continue

        if in_bullet_section and _is_section_header(line_upper):
            if "BULLET POINT" not in line_upper:
                in_bullet_section = False
                if buffer:
                    bullets.append(_parse_bullet_line(buffer))
                    buffer = ""
                continue

        if not in_bullet_section:
            continue

        if re.match(r"^[*_\s]*(Unique Bullet|Note:|Label:|Unique BP)", line, re.I):
            continue
        if re.match(r"^[^(]+:\s*$", line) and not _extract_asins(line):
            continue

        line_asins = _extract_asins(line)
        has_asin_suffix = bool(line_asins)
        has_dash = bool(re.match(r"^[-•]\s+", line))
        is_new_bullet = has_dash or has_asin_suffix

        if is_new_bullet:
            if buffer:
                bullets.append(_parse_bullet_line(buffer))
                buffer = ""
            buffer = re.sub(r"^[-•]\s+", "", line)
        elif buffer:
            looks_like_new = len(line) > 40 and bool(re.match(r"^[A-Z0-9][^:]{3,80}:", line))
            if looks_like_new:
                bullets.append(_parse_bullet_line(buffer))
                buffer = line
            else:
                buffer += " " + line
        else:
            if len(line) > 40:
                buffer = line

    if buffer:
        bullets.append(_parse_bullet_line(buffer))

    return [b for b in bullets if b["text"] and len(b["text"]) > 10]


def _parse_description(lines: list[str]) -> str:
    """
    The real product description lives in the table cell labeled exactly
    "Brand Description" (under the BRAND DESCRIPTION / BACKEND DESCRIPTION
    headers) — not the EBC/A+ module content, which is separate marketing
    copy for the designer, not the description field itself. The value can
    span several paragraphs within that one cell, so this collects every
    line after the label until end-of-document or the closing boilerplate.
    """
    label_idx = None
    for i, line in enumerate(lines):
        if line.strip().lower() == "brand description":
            label_idx = i  # keep the LAST match — the section header "BRAND DESCRIPTION"
                           # appears earlier and lowercases to the same string as the
                           # actual table label that follows it, so take the final one

    if label_idx is None:
        return ""

    stop_markers = ("ONE MORE THING", "PLEASE LEAVE US A COMMENT")
    parts = []
    for line in lines[label_idx + 1:]:
        upper = line.upper()
        if any(marker in upper for marker in stop_markers):
            break
        parts.append(line)

    return " ".join(parts).strip()


def _parse_keywords(lines: list[str]) -> str:
    in_seo = False
    kw_lines = []
    for line in lines:
        upper = line.upper()
        entered = False
        for kw in SEO_HEADERS:
            if kw in upper:
                in_seo = True
                entered = True
                break
        if entered:
            continue
        if not in_seo:
            continue
        stopped = False
        for kw in SEO_STOP_HEADERS:
            if kw in upper:
                in_seo = False
                stopped = True
                break
        if stopped:
            continue
        if any(m in upper for m in SEO_MARKERS):
            continue
        word_count = line.count(" ") + 1
        has_sentence = bool(re.search(r"[.!?]\s+[A-Z]", line))
        if len(line) > 30 and word_count > 5 and not has_sentence:
            kw_lines.append(line)
    return " ".join(kw_lines).strip()


def parse(file_path: str, title_mode: str = "auto") -> dict:
    """
    Parse a .docx copy template.

    Returns:
      {'needs_title_choice': bool, 'title_modes': [...], 'asins': {ASIN: {...}}}
    'asins' values include title, sku, brand, bullet_1..5, description, keywords.
    """
    lines = extract_lines(file_path)

    has_test = _has_section(lines, "TEST TITLE")
    has_current = _has_section(lines, "CURRENT TITLE")
    needs_choice = has_test and has_current

    if needs_choice and title_mode == "auto":
        return {"needs_title_choice": True, "title_modes": ["test", "current"], "asins": {}}

    asins = _parse_titles(lines, title_mode, has_test, has_current)
    bullets = _parse_bullets(lines)

    if not asins:
        raise ValueError(
            "No ASINs found in the document. Make sure titles are followed by an "
            "ASIN in parentheses, e.g. \"Title text (B0XXXXXXXX)\"."
        )

    all_asin_codes = list(asins.keys())
    for bullet in bullets:
        targets = bullet["asins"] if bullet["asins"] else all_asin_codes
        for asin in targets:
            if asin not in asins:
                continue
            for i in range(1, 6):
                key = f"bullet_{i}"
                if not asins[asin][key]:
                    asins[asin][key] = bullet["text"]
                    break

    description = _parse_description(lines)
    keywords = _parse_keywords(lines)
    for asin in asins:
        asins[asin]["description"] = description
        asins[asin]["keywords"] = keywords

    return {
        "needs_title_choice": False,
        "title_modes": ["test", "current"] if needs_choice else [],
        "asins": asins,
    }
